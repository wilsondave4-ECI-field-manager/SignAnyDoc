import base64, io, json, random, socket, threading, time, webbrowser, sys
from http.server import ThreadingHTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import fitz
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Inches
import qrcode

try:
    import win32com.client
    WORD_COM_AVAILABLE=True
except Exception:
    WORD_COM_AVAILABLE=False

def resource_path(name):
    base = Path(getattr(sys, '_MEIPASS', Path(__file__).resolve().parent))
    return base / name

APP_NAME='Doc Sign'
PORT=8765
PHONE_HTML=resource_path('phone.html').read_text(encoding='utf-8')
STATE={'code':''.join(random.choice('ABCDEFGHJKLMNPQRSTUVWXYZ23456789') for _ in range(6)),
       'pin':str(random.randint(1000,9999)),'paired':False,'signature':None}
LOCK=threading.Lock()

def local_ip():
    try:
        s=socket.socket(socket.AF_INET,socket.SOCK_DGRAM);s.connect(('8.8.8.8',80))
        ip=s.getsockname()[0];s.close();return ip
    except Exception:
        try:return socket.gethostbyname(socket.gethostname())
        except:return '127.0.0.1'

class Handler(BaseHTTPRequestHandler):
    def send_json(self,data,status=200):
        b=json.dumps(data).encode()
        self.send_response(status);self.send_header('Content-Type','application/json')
        self.send_header('Cache-Control','no-store');self.send_header('Content-Length',str(len(b)))
        self.end_headers();self.wfile.write(b)
    def body(self):
        n=int(self.headers.get('Content-Length','0'))
        return json.loads(self.rfile.read(n) or b'{}')
    def do_GET(self):
        p=urlparse(self.path).path
        if p in ('/','/phone'):
            b=PHONE_HTML.encode();self.send_response(200);self.send_header('Content-Type','text/html; charset=utf-8')
            self.send_header('Content-Length',str(len(b)));self.end_headers();self.wfile.write(b);return
        if p=='/api/state':
            with LOCK:self.send_json(dict(STATE));return
        self.send_error(404)
    def do_POST(self):
        p=urlparse(self.path).path
        if p=='/api/pair':
            d=self.body()
            with LOCK:
                ok=d.get('code','').upper()==STATE['code'] and str(d.get('pin',''))==STATE['pin']
                if ok:STATE['paired']=True
            return self.send_json({'ok':ok},200 if ok else 403)
        if p=='/api/signature':
            d=self.body()
            with LOCK:
                ok=d.get('code','').upper()==STATE['code'] and str(d.get('pin',''))==STATE['pin']
                sig=d.get('signature','')
                if ok and sig.startswith('data:image/png;base64,'):STATE['signature']=sig
                else:ok=False
            return self.send_json({'ok':ok},200 if ok else 403)
        self.send_error(404)
    def log_message(self,fmt,*args): pass

class App:
    def __init__(self,root):
        self.root=root;root.title('Doc Sign');root.geometry('1180x820');root.minsize(980,700)
        self.doc=None;self.path=None;self.doc_type=None;self.page_index=0
        self.render_image=None;self.tk_image=None;self.signature_pil=None;self.signature_tk=None
        self.sig_x=.58;self.sig_y=.72;self.sig_w=.24
        self.server=None
        self.build_ui();self.start_server();self.refresh_pairing();self.poll_signature()

    def build_ui(self):
        top=ttk.Frame(self.root,padding=12);top.pack(fill='x')
        ttk.Label(top,text='Doc Sign - Local Document Signing',font=('Segoe UI',18,'bold')).pack(side='left')
        ttk.Button(top,text='Open PDF',command=self.open_pdf).pack(side='right',padx=4)
        ttk.Button(top,text='Open Word',command=self.open_word).pack(side='right',padx=4)

        body=ttk.Panedwindow(self.root,orient='horizontal');body.pack(fill='both',expand=True,padx=12,pady=(0,12))
        left=ttk.Frame(body,padding=8);right=ttk.Frame(body,padding=8,width=330);body.add(left,weight=4);body.add(right,weight=1)

        toolbar=ttk.Frame(left);toolbar.pack(fill='x')
        self.file_label=ttk.Label(toolbar,text='No document loaded');self.file_label.pack(side='left')
        self.page_label=ttk.Label(toolbar,text='');self.page_label.pack(side='right')
        ttk.Button(toolbar,text='<',width=4,command=self.prev_page).pack(side='right',padx=2)
        ttk.Button(toolbar,text='>',width=4,command=self.next_page).pack(side='right',padx=2)

        self.canvas=tk.Canvas(left,bg='#e8edf2',highlightthickness=1,highlightbackground='#ccd6de')
        self.canvas.pack(fill='both',expand=True,pady=8)
        self.canvas.bind('<Configure>',lambda e:self.render_current())
        self.canvas.bind('<Button-1>',self.place_signature)

        actions=ttk.Frame(left);actions.pack(fill='x')
        ttk.Label(actions,text='Signature size').pack(side='left')
        ttk.Button(actions,text='-',width=4,command=lambda:self.resize_sig(-.03)).pack(side='left',padx=3)
        ttk.Button(actions,text='+',width=4,command=lambda:self.resize_sig(.03)).pack(side='left')
        self.save_btn=ttk.Button(actions,text='Save signed document',command=self.save_signed,state='disabled')
        self.save_btn.pack(side='right')
        self.word_cursor_btn=ttk.Button(actions,text='Insert at Word cursor',command=self.insert_at_word_cursor,state='disabled')
        self.word_cursor_btn.pack(side='right',padx=8)

        pair=ttk.LabelFrame(right,text='Phone pairing',padding=14);pair.pack(fill='x')
        self.code_label=ttk.Label(pair,text='',font=('Segoe UI',22,'bold'));self.code_label.pack()
        self.pin_label=ttk.Label(pair,text='',font=('Segoe UI',18,'bold'));self.pin_label.pack(pady=(0,8))
        self.qr_label=ttk.Label(pair);self.qr_label.pack(pady=6)
        self.url_box=tk.Text(pair,height=4,width=34,wrap='word');self.url_box.pack(fill='x');self.url_box.configure(state='disabled')
        ttk.Button(pair,text='Open phone page on this PC',command=self.open_phone_local).pack(fill='x',pady=(8,0))
        ttk.Button(pair,text='New session',command=self.new_session).pack(fill='x',pady=5)
        self.pair_status=ttk.Label(pair,text='Waiting for phone');self.pair_status.pack(pady=5)

        sig=ttk.LabelFrame(right,text='Received signature',padding=14);sig.pack(fill='x',pady=12)
        self.sig_preview=ttk.Label(sig,text='No signature received');self.sig_preview.pack(pady=8)
        ttk.Button(sig,text='Clear signature',command=self.clear_signature).pack(fill='x')

        note=ttk.LabelFrame(right,text='Security',padding=12);note.pack(fill='x')
        ttk.Label(note,text='Local network only\nSession code + PIN\nSignature kept in memory\nOriginal document is not overwritten',
                  justify='left').pack(anchor='w')

        self.status=ttk.Label(self.root,text='Ready.',anchor='w',padding=(12,4));self.status.pack(fill='x')

    def start_server(self):
        self.server=ThreadingHTTPServer(('0.0.0.0',PORT),Handler)
        threading.Thread(target=self.server.serve_forever,daemon=True).start()

    def phone_url(self):
        return f"http://{local_ip()}:{PORT}/phone?code={STATE['code']}&pin={STATE['pin']}"

    def refresh_pairing(self):
        self.code_label.config(text=f"SESSION  {STATE['code']}")
        self.pin_label.config(text=f"PIN  {STATE['pin']}")
        url=self.phone_url()
        self.url_box.configure(state='normal');self.url_box.delete('1.0','end');self.url_box.insert('1.0',url);self.url_box.configure(state='disabled')
        qr=qrcode.make(url).resize((210,210))
        self.qr_tk=ImageTk.PhotoImage(qr);self.qr_label.config(image=self.qr_tk)
        self.pair_status.config(text='Phone connected' if STATE['paired'] else 'Waiting for phone')

    def new_session(self):
        with LOCK:
            STATE['code']=''.join(random.choice('ABCDEFGHJKLMNPQRSTUVWXYZ23456789') for _ in range(6))
            STATE['pin']=str(random.randint(1000,9999));STATE['paired']=False;STATE['signature']=None
        self.signature_pil=None;self.sig_preview.config(image='',text='No signature received');self.save_btn.config(state='disabled');self.word_cursor_btn.config(state='disabled')
        self.refresh_pairing();self.render_current();self.status.config(text='New local signing session created.')

    def open_phone_local(self):
        webbrowser.open(f"http://127.0.0.1:{PORT}/phone?code={STATE['code']}&pin={STATE['pin']}")

    def poll_signature(self):
        self.refresh_pairing()
        with LOCK:sig=STATE['signature']
        if sig and self.signature_pil is None:
            try:
                raw=base64.b64decode(sig.split(',',1)[1]);im=Image.open(io.BytesIO(raw)).convert('RGBA')
                alpha=im.getchannel('A');bbox=alpha.getbbox()
                if bbox:im=im.crop(bbox)
                self.signature_pil=im
                prev=im.copy();prev.thumbnail((250,100));self.sig_prev_tk=ImageTk.PhotoImage(prev)
                self.sig_preview.config(image=self.sig_prev_tk,text='')
                if self.doc_type:self.save_btn.config(state='normal')
                if self.doc_type=='docx' and WORD_COM_AVAILABLE:self.word_cursor_btn.config(state='normal')
                self.render_current();self.status.config(text='Signature received directly from phone.')
            except Exception as e:self.status.config(text=f'Could not read signature: {e}')
        self.root.after(600,self.poll_signature)

    def clear_signature(self):
        with LOCK:STATE['signature']=None
        self.signature_pil=None;self.sig_preview.config(image='',text='No signature received');self.save_btn.config(state='disabled');self.word_cursor_btn.config(state='disabled')
        self.render_current();self.status.config(text='Signature cleared.')

    def open_pdf(self):
        p=filedialog.askopenfilename(filetypes=[('PDF documents','*.pdf')])
        if not p:return
        try:
            if self.doc:self.doc.close()
            self.doc=fitz.open(p);self.path=Path(p);self.doc_type='pdf';self.page_index=0
            self.file_label.config(text=self.path.name);self.status.config(text='PDF loaded locally. Click the page to position the signature.')
            if self.signature_pil:self.save_btn.config(state='normal')
            self.word_cursor_btn.config(state='disabled')
            self.render_current()
        except Exception as e:messagebox.showerror('Open PDF',str(e))

    def open_word(self):
        p=filedialog.askopenfilename(filetypes=[('Word documents','*.docx')])
        if not p:return
        self.path=Path(p);self.doc_type='docx';self.doc=None;self.page_index=0
        self.file_label.config(text=self.path.name);self.page_label.config(text='Word document')
        self.canvas.delete('all')
        w=max(self.canvas.winfo_width(),600);h=max(self.canvas.winfo_height(),500)
        self.canvas.create_rectangle(80,40,w-80,h-40,fill='white',outline='#ccd6de')
        self.canvas.create_text(w/2,120,text='Microsoft Word (.docx)',font=('Segoe UI',22,'bold'),fill='#10243d')
        self.canvas.create_text(w/2,180,text='This local build appends the received signature near the end of the Word document.',width=w-240,font=('Segoe UI',13),fill='#5d7081')
        if self.signature_pil:
            self.canvas.create_text(w/2,250,text='Signature received - ready to save or insert at the current Word cursor.',font=('Segoe UI',14,'bold'),fill='#177a72')
            self.save_btn.config(state='normal')
            self.word_cursor_btn.config(state='normal' if WORD_COM_AVAILABLE else 'disabled')
        else:
            self.save_btn.config(state='disabled')
            self.word_cursor_btn.config(state='disabled')
        note='Word document loaded locally. For exact placement, open the document in Microsoft Word, click the target position, then use Insert at Word cursor.' if WORD_COM_AVAILABLE else 'Word loaded. Native cursor placement needs Microsoft Word plus the Windows COM component.'
        self.status.config(text=note)

    def prev_page(self):
        if self.doc_type=='pdf' and self.doc and self.page_index>0:self.page_index-=1;self.render_current()
    def next_page(self):
        if self.doc_type=='pdf' and self.doc and self.page_index<len(self.doc)-1:self.page_index+=1;self.render_current()

    def render_current(self):
        if self.doc_type!='pdf' or not self.doc:return
        try:
            page=self.doc[self.page_index]
            cw=max(self.canvas.winfo_width()-30,300);ch=max(self.canvas.winfo_height()-30,300)
            r=page.rect;scale=min(cw/r.width,ch/r.height)
            pix=page.get_pixmap(matrix=fitz.Matrix(scale,scale),alpha=False)
            im=Image.frombytes('RGB',[pix.width,pix.height],pix.samples)
            self.render_image=im;self.tk_image=ImageTk.PhotoImage(im)
            self.canvas.delete('all');x=(self.canvas.winfo_width()-pix.width)//2;y=(self.canvas.winfo_height()-pix.height)//2
            self.page_origin=(x,y,pix.width,pix.height)
            self.canvas.create_image(x,y,anchor='nw',image=self.tk_image)
            if self.signature_pil:
                sw=max(70,int(pix.width*self.sig_w));ratio=self.signature_pil.height/max(self.signature_pil.width,1);sh=max(24,int(sw*ratio))
                s=self.signature_pil.resize((sw,sh),Image.Resampling.LANCZOS);self.signature_tk=ImageTk.PhotoImage(s)
                sx=x+int(self.sig_x*pix.width)-sw//2;sy=y+int(self.sig_y*pix.height)-sh//2
                self.canvas.create_image(sx,sy,anchor='nw',image=self.signature_tk,tags='sig')
            self.page_label.config(text=f'Page {self.page_index+1} of {len(self.doc)}')
        except Exception as e:self.status.config(text=f'Preview error: {e}')

    def place_signature(self,event):
        if self.doc_type!='pdf' or not self.doc or not self.signature_pil:return
        x,y,w,h=self.page_origin
        if x<=event.x<=x+w and y<=event.y<=y+h:
            self.sig_x=(event.x-x)/w;self.sig_y=(event.y-y)/h;self.render_current()

    def resize_sig(self,d):
        self.sig_w=max(.08,min(.5,self.sig_w+d));self.render_current()

    def insert_at_word_cursor(self):
        if not self.path or self.doc_type!='docx' or not self.signature_pil:
            messagebox.showinfo('Word placement','Open a Word document and receive a signature first.')
            return
        if not WORD_COM_AVAILABLE:
            messagebox.showerror('Word placement','Microsoft Word cursor placement is not available in this build.')
            return
        try:
            temp = Path.home()/'AppData'/'Local'/'Doc Sign'
            temp.mkdir(parents=True,exist_ok=True)
            png_path = temp/'current_signature.png'
            self.signature_pil.save(png_path,format='PNG')

            word = win32com.client.GetActiveObject('Word.Application')
            doc = word.ActiveDocument
            active_path = Path(doc.FullName).resolve()
            target_path = self.path.resolve()
            if active_path != target_path:
                answer=messagebox.askyesno('Word document check','The active Microsoft Word document is not the same file loaded in Doc Sign.\n\n'+f'Loaded in Doc Sign:\n{target_path.name}\n\n'+f'Active in Word:\n{active_path.name}\n\n'+'Insert the signature into the active Word document anyway?')
                if not answer:return

            sel = word.Selection
            inline = sel.InlineShapes.AddPicture(str(png_path),False,True)
            inline.LockAspectRatio = True
            inline.Width = 180
            self.status.config(text='Signature inserted at the current Microsoft Word cursor position.')
            messagebox.showinfo('Word placement','Signature inserted at the current Word cursor.\n\nUse Word Save or Save As to keep the signed document.')
        except Exception as e:
            messagebox.showerror('Word placement','Could not insert the signature into Microsoft Word.\n\nMake sure Microsoft Word is open, the document is active, and the cursor is at the required position.\n\n'+str(e))

    def save_signed(self):
        if not self.path or not self.signature_pil:return
        if self.doc_type=='pdf':self.save_pdf()
        elif self.doc_type=='docx':self.save_word()

    def save_pdf(self):
        default=self.path.stem+'_Signed.pdf'
        out=filedialog.asksaveasfilename(defaultextension='.pdf',initialfile=default,filetypes=[('PDF','*.pdf')])
        if not out:return
        try:
            doc=fitz.open(str(self.path));page=doc[self.page_index];rect=page.rect
            sw=rect.width*self.sig_w;ratio=self.signature_pil.height/max(self.signature_pil.width,1);sh=sw*ratio
            cx=self.sig_x*rect.width;cy=self.sig_y*rect.height
            target=fitz.Rect(cx-sw/2,cy-sh/2,cx+sw/2,cy+sh/2)
            bio=io.BytesIO();self.signature_pil.save(bio,format='PNG')
            page.insert_image(target,stream=bio.getvalue(),overlay=True,keep_proportion=True)
            doc.save(out);doc.close();self.status.config(text=f'Signed PDF saved: {out}')
            messagebox.showinfo('Saved','Signed PDF saved.\n\nThe original document was not changed.')
        except Exception as e:messagebox.showerror('Save PDF',str(e))

    def save_word(self):
        default=self.path.stem+'_Signed.docx'
        out=filedialog.asksaveasfilename(defaultextension='.docx',initialfile=default,filetypes=[('Word document','*.docx')])
        if not out:return
        try:
            doc=Document(str(self.path));doc.add_paragraph()
            p=doc.add_paragraph();p.add_run('Signature:')
            bio=io.BytesIO();self.signature_pil.save(bio,format='PNG');bio.seek(0)
            p.add_run().add_picture(bio,width=Inches(2.5))
            doc.save(out);self.status.config(text=f'Signed Word document saved: {out}')
            messagebox.showinfo('Saved','Signed Word document saved.\n\nThe original document was not changed.')
        except Exception as e:messagebox.showerror('Save Word',str(e))

def main():
    root=tk.Tk()
    try:
        style=ttk.Style();style.theme_use('vista')
    except:pass
    app=App(root)
    root.protocol('WM_DELETE_WINDOW',lambda:(app.server.shutdown() if app.server else None,root.destroy()))
    root.mainloop()

if __name__=='__main__':main()
